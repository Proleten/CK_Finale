import pandas as pd
from docx import Document
from tkinter import Tk, Label, Entry, Button, messagebox
from datetime import datetime

def fill_word_template(org_name):
    excel_path = r"C:\Users\fullo\Desktop\Портфель\Проект ЦК\FORTEST.xlsx"

    try:
        df = pd.read_excel(excel_path)

        row = df[df['NAME'] == org_name]

        if row.empty:
            messagebox.showerror("Брехня", f"Организация '{org_name}' не найдена в исходной таблице.")
            return

        data = row.iloc[0].to_dict()

        template_path = r"C:\Users\fullo\Desktop\Портфель\Проект ЦК\finaltest.docx"
        doc = Document(template_path)

        replacements = {
            'SUD': data.get('SUD', ''),
            'NAME': data.get('NAME', ''),
            'ADDRESS': data.get('ADDRESS', ''),
            'ID': data.get('ID', ''),
            'REG': data.get('REG', ''),
            'DEBT': data.get('DEBT', ''),
            'PENI': data.get('PENI', ''),
            'PERIOD': data.get('PERIOD', ''),
            'GP': data.get('GP', ''),
            'PRETENZIA': data.get('PRETENZIA', ''),
            'DATE': datetime.now().strftime('%d.%m.%Y')
        }

        def replace_text_in_paragraph(paragraph, replacements):
            for key, value in replacements.items():
                placeholder = f'{{{key}}}'
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

        def replace_text_in_table(table, replacements):
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        placeholder = f'{{{key}}}'
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))

        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)

        for table in doc.tables:
            replace_text_in_table(table, replacements)

        output_path = f"Исковое заявление {data['NAME']}.docx"
        doc.save(output_path)

        messagebox.showinfo("Ну сейчас начнется...", f"Документ успешно создан: {output_path}")

    except Exception as e:
        messagebox.showerror("Что-то не так...", str(e))


def on_generate_click():
    org_name = entry.get()
    if org_name:
        fill_word_template(org_name)
    else:
        messagebox.showwarning("Так не пойдёт...", "Пожалуйста, укажите название организации.")


root = Tk()
root.title("Генератор исков")
root.geometry("400x300")
root.config(bg="wheat3")

Label(root, text="Укажите название организации:", justify='center', background="wheat", font="Cambria 12").pack(pady=25)
entry = Entry(root, width=50)
entry.pack(pady=25)

Button(root, text="Сгенерировать иск", command=on_generate_click, background="light green", font="Cambria 12").pack(pady=25)

root.mainloop()
